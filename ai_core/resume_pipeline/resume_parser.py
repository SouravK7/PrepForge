"""
Resume parser.

Extracts raw text from PDF, DOCX, and TXT resume files.
Detects document structure and section boundaries.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Optional


class ResumeParser:
    """
    Extracts and structures raw text from resume files.

    Supports PDF, DOCX, and TXT formats.
    Detects section boundaries for structured extraction.
    """

    # Section header patterns
    SECTION_PATTERNS: dict[str, re.Pattern] = {
        "education": re.compile(
            r"(education|academic|qualification|degree|university|college)",
            re.IGNORECASE,
        ),
        "experience": re.compile(
            r"(experience|employment|work history|career|professional background|internship)",
            re.IGNORECASE,
        ),
        "skills": re.compile(
            r"(skills|technical skills|technologies|competencies|expertise|tools)",
            re.IGNORECASE,
        ),
        "projects": re.compile(
            r"(projects|portfolio|personal projects|academic projects|side projects)",
            re.IGNORECASE,
        ),
        "certifications": re.compile(
            r"(certifications|certificates|courses|training|achievements|awards)",
            re.IGNORECASE,
        ),
        "summary": re.compile(
            r"(summary|profile|objective|about|overview|introduction)",
            re.IGNORECASE,
        ),
    }

    def parse_file(self, file_path: str | Path) -> dict:
        """
        Parse a resume file and extract structured text.

        Args:
            file_path: Path to the resume file.

        Returns:
            Dict with raw_text, sections, file_type, char_count, word_count.

        Raises:
            ValueError: If file format is unsupported.
            FileNotFoundError: If file does not exist.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            raw_text = self._parse_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            raw_text = self._parse_docx(file_path)
        elif suffix == ".txt":
            raw_text = self._parse_txt(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {suffix}. "
                f"Supported: .pdf, .docx, .doc, .txt"
            )

        sections = self._detect_sections(raw_text)

        return {
            "raw_text": raw_text,
            "sections": sections,
            "file_type": suffix.lstrip("."),
            "char_count": len(raw_text),
            "word_count": len(raw_text.split()),
        }

    def parse_text(self, raw_text: str) -> dict:
        """
        Parse already-extracted raw text.

        Args:
            raw_text: Plain text content of resume.

        Returns:
            Dict with raw_text and detected sections.
        """
        sections = self._detect_sections(raw_text)
        return {
            "raw_text": raw_text,
            "sections": sections,
            "file_type": "text",
            "char_count": len(raw_text),
            "word_count": len(raw_text.split()),
        }

    def parse_bytes(self, file_bytes: bytes, file_name: str) -> dict:
        """
        Parse resume from bytes (e.g. from Streamlit file_uploader).

        Args:
            file_bytes: File content as bytes.
            file_name: Original file name with extension.

        Returns:
            Parsed resume dict.

        Raises:
            ValueError: If file format unsupported.
        """
        suffix = Path(file_name).suffix.lower()

        if suffix == ".pdf":
            raw_text = self._parse_pdf_bytes(file_bytes)
        elif suffix in (".docx", ".doc"):
            raw_text = self._parse_docx_bytes(file_bytes)
        elif suffix == ".txt":
            raw_text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        sections = self._detect_sections(raw_text)

        return {
            "raw_text": raw_text,
            "sections": sections,
            "file_type": suffix.lstrip("."),
            "char_count": len(raw_text),
            "word_count": len(raw_text.split()),
        }

    # ── Private parsers ────────────────────────────────────────

    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError(
                "pdfplumber not installed. Run: pip install pdfplumber"
            )
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")

    def _parse_pdf_bytes(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("pdfplumber not installed.")
        except Exception as e:
            raise ValueError(f"Failed to parse PDF bytes: {e}")

    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = [
                para.text for para in doc.paragraphs if para.text.strip()
            ]
            return "\n".join(paragraphs)
        except ImportError:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")

    def _parse_docx_bytes(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [
                para.text for para in doc.paragraphs if para.text.strip()
            ]
            return "\n".join(paragraphs)
        except ImportError:
            raise ImportError("python-docx not installed.")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX bytes: {e}")

    def _parse_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Failed to read TXT file: {e}")

    # Characters that indicate a line is content, not a section header
    _CONTENT_CHARS: frozenset[str] = frozenset({":", "—", "–", "@", "|"})

    def _detect_sections(self, text: str) -> dict[str, str]:
        """
        Detect and extract resume sections from raw text.

        Splits text at section boundaries detected by header patterns.
        Only short, plain lines are treated as section headers.
        Lines with colons, em-dashes, or parenthesised years are always
        treated as content even if they contain a section keyword.

        Args:
            text: Raw resume text.

        Returns:
            Dict of section_name to section_text.
        """
        import re as _re

        _year_in_parens = _re.compile(r"\(\d{4}\)")
        _date_range = _re.compile(r"\d{4}\s*[-–—]\s*\d{4}")

        lines = text.split("\n")
        sections: dict[str, list[str]] = {
            name: [] for name in self.SECTION_PATTERNS
        }
        sections["other"] = []

        current_section = "other"

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # A line qualifies as a section header only if:
            #   - it is short (< 40 chars)
            #   - it contains no content-indicator characters (colon, dash…)
            #   - it has no parenthesised year like "(2020)"
            #   - it has no date range like "2020-2024"
            is_header_candidate = (
                len(stripped) < 40
                and not any(ch in stripped for ch in self._CONTENT_CHARS)
                and not _year_in_parens.search(stripped)
                and not _date_range.search(stripped)
            )

            matched_section = None
            if is_header_candidate:
                for section_name, pattern in self.SECTION_PATTERNS.items():
                    if pattern.search(stripped):
                        matched_section = section_name
                        break

            if matched_section:
                current_section = matched_section
            else:
                sections[current_section].append(stripped)

        return {
            name: "\n".join(lines_list)
            for name, lines_list in sections.items()
            if lines_list
        }
